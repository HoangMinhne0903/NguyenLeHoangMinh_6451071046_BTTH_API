from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
OUTPUT = ROOT / "Bao_cao_da_dien_noi_dung_Facebook_Page_API.docx"

ACCENT = RGBColor(46, 116, 181)
GRAY = RGBColor(90, 90, 90)
BOX_FILL = "F7F7F7"


def set_font(run, size=11, bold=False, italic=False, color=None):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def set_document_defaults(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)

    for name, size in [("Heading 1", 14), ("Heading 2", 12)]:
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = ACCENT


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_para(doc: Document, text: str, bold=False, italic=False, center=False, color=None, after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, bold=bold, italic=italic, color=color)
    return p


def add_section(doc: Document, title: str):
    doc.add_paragraph(title, style="Heading 1")


def add_subsection(doc: Document, title: str):
    doc.add_paragraph(title, style="Heading 2")


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_font(run)


def add_figure_box(doc: Document, title: str, note: str = ""):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, BOX_FILL)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    r1 = p.add_run(title + "\n")
    set_font(r1, bold=True)
    r2 = p.add_run("[Chèn ảnh minh chứng tại đây]")
    set_font(r2, italic=True, color=GRAY)
    if note:
        r3 = p.add_run("\n" + note)
        set_font(r3, size=10, color=GRAY)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(title)
    set_font(r, size=10, italic=True)
    cap.paragraph_format.space_after = Pt(6)


def add_info_table(doc: Document, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, (left, right) in enumerate(rows):
        shade_cell(table.cell(index, 0), "F2F4F7")
        c1 = table.cell(index, 0)
        c2 = table.cell(index, 1)
        c1.text = ""
        c2.text = ""
        p1 = c1.paragraphs[0]
        p2 = c2.paragraphs[0]
        set_font(p1.add_run(left), bold=True)
        set_font(p2.add_run(right))
        c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def build():
    doc = Document()
    set_document_defaults(doc)

    add_para(doc, "Họ và tên: [Điền họ và tên]", bold=True, after=2)
    add_para(doc, "MSSV: [Điền mã số sinh viên]", bold=True, after=10)

    add_section(doc, "1. Webhook")
    add_para(
        doc,
        "Trong hệ thống này, webhook-service là service đầu tiên tiếp nhận sự kiện từ Facebook Page. "
        "Service được cấu hình chạy ở port 3001 và expose endpoint /webhook để Facebook gửi sự kiện về. "
        "Ngoài ra, service còn hỗ trợ bước xác thực ban đầu thông qua Verify Token để Facebook xác nhận callback URL hợp lệ."
    )
    add_para(
        doc,
        "Trong quá trình triển khai, em sử dụng ngrok để public port 3001 ra Internet vì Facebook không thể gọi trực tiếp tới localhost. "
        "Sau khi chạy ngrok, em lấy đường dẫn public dạng https://<ngrok-domain>/webhook và khai báo vào phần Webhooks trên Meta for Developers. "
        "Verify Token trong Facebook Developer phải trùng với giá trị VerifyToken cấu hình trong appsettings của webhook-service."
    )
    add_bullet(doc, "webhook-service chạy tại port 3001.")
    add_bullet(doc, "Route webhook dùng để verify và nhận event là /webhook.")
    add_bullet(doc, "Sự kiện nhận được sẽ được publish vào Kafka topic raw_events.")
    add_bullet(doc, "Hệ thống đã verify webhook thành công trên Meta for Developers.")
    add_info_table(
        doc,
        [
            ("Port", "3001"),
            ("Route", "/webhook"),
            ("Object subscribe", "Page"),
            ("Field subscribe", "feed"),
            ("Vai trò", "Nhận event từ Facebook và đẩy vào Kafka"),
        ],
    )
    add_figure_box(doc, "Hình 1. Ảnh webhook-service đang chạy trên port 3001")
    add_figure_box(doc, "Hình 2. Ảnh cấu hình webhook trên Facebook Developer", "Nên chụp phần Callback URL, Verify Token và field feed.")
    add_figure_box(doc, "Hình 3. Ảnh comment thật trên bài viết của page")
    add_figure_box(doc, "Hình 4. Ảnh Kafka UI nhận dữ liệu trong topic raw_events")

    add_section(doc, "2. Core Service")
    add_para(
        doc,
        "Core-service là service chịu trách nhiệm xử lý nghiệp vụ chính của hệ thống. "
        "Sau khi consume dữ liệu từ topic raw_events, service sẽ phân tích loại sự kiện, tách nội dung comment hoặc message, "
        "đánh giá đây là bình luận bình thường, câu hỏi hỏi giá, khiếu nại, spam hay nội dung tiêu cực. "
        "Từ kết quả đó, core-service sẽ sinh command phù hợp và publish sang topic reply_commands để backend-api thực thi."
    )
    add_para(
        doc,
        "Trong quá trình hoàn thiện bài, em có bổ sung logic để tránh loop khi page tự phản hồi chính comment của mình. "
        "Cụ thể, các event do chính Page sinh ra sẽ bị bỏ qua trước khi đưa lại vào pipeline. "
        "Ngoài ra, service cũng có rule xử lý spam, duplicate comment, nội dung toxic nặng và blacklist nội bộ theo số lần spam."
    )
    add_bullet(doc, "core-service chạy tại port 3002.")
    add_bullet(doc, "Input: Kafka topic raw_events.")
    add_bullet(doc, "Output: Kafka topic reply_commands.")
    add_bullet(doc, "Có xử lý rule spam, toxic comment, negative comment và self-loop.")
    add_info_table(
        doc,
        [
            ("Port", "3002"),
            ("Input topic", "raw_events"),
            ("Output topic", "reply_commands"),
            ("Vai trò", "Phân tích nội dung và sinh command phản hồi"),
        ],
    )
    add_figure_box(doc, "Hình 5. Ảnh core-service đang chạy")
    add_figure_box(doc, "Hình 6. Ảnh bình luận tích cực hoặc hỏi giá để test phân loại nội dung")
    add_figure_box(doc, "Hình 7. Ảnh Kafka UI sau khi core-service xử lý xong")
    add_figure_box(doc, "Hình 8. Ảnh trường hợp spam hoặc toxic comment")

    add_section(doc, "3. Backend API")
    add_para(
        doc,
        "Backend-api là service duy nhất trong hệ thống được phép gọi Facebook Graph API. "
        "Service này consume dữ liệu từ topic reply_commands và send_retry, sau đó thực thi hành động thực tế như reply comment hoặc hide comment trên Facebook Page. "
        "Ngoài ra, backend-api cũng lưu dấu vết command đã xử lý để phục vụ idempotency và lưu thông tin lỗi vào database khi có command thất bại."
    )
    add_para(
        doc,
        "Trong bài làm này, em cấu hình SQL Server chạy local qua Docker và kết nối backend-api tới database ApiFacebookDb. "
        "Các bảng chính được tạo gồm EventTrackings, FailedCommands và IdempotencyKeys. "
        "Khi Page Access Token hết hạn hoặc comment id không hợp lệ, backend-api sẽ đẩy command lỗi sang topic send_failed để retry-service tiếp tục xử lý."
    )
    add_bullet(doc, "backend-api chạy tại port 3000.")
    add_bullet(doc, "Consume topic reply_commands và send_retry.")
    add_bullet(doc, "Kết nối SQL Server tại localhost:1435.")
    add_bullet(doc, "Gọi Facebook Graph API để reply hoặc hide comment.")
    add_info_table(
        doc,
        [
            ("Port", "3000"),
            ("Input topic", "reply_commands, send_retry"),
            ("Database", "ApiFacebookDb trên SQL Server"),
            ("Vai trò", "Thực thi command thật trên Facebook và lưu dữ liệu lỗi"),
        ],
    )
    add_figure_box(doc, "Hình 9. Ảnh backend-api đang chạy")
    add_figure_box(doc, "Hình 10. Ảnh health endpoint hoặc log consume command")
    add_figure_box(doc, "Hình 11. Ảnh topic reply_commands trong Kafka UI")
    add_figure_box(doc, "Hình 12. Ảnh dữ liệu trong database hoặc bảng failed_commands")

    add_section(doc, "4. Retry Service")
    add_para(
        doc,
        "Retry-service được dùng để xử lý các command thất bại khi backend-api không gửi được yêu cầu lên Facebook Graph API. "
        "Service này consume topic send_failed, chờ theo thời gian backoff tăng dần, sau đó republish command vào send_retry để backend-api thử lại. "
        "Nếu command vượt quá số lần retry cho phép thì hệ thống sẽ đưa command đó vào dead_letter để tránh lặp vô hạn."
    )
    add_para(
        doc,
        "Việc tách riêng retry-service giúp hệ thống có khả năng phục hồi tốt hơn khi gặp lỗi tạm thời như token sai, mạng chậm hoặc API bên ngoài phản hồi thất bại. "
        "Trong báo cáo, phần này nên minh họa rõ 3 topic liên quan là send_failed, send_retry và dead_letter."
    )
    add_bullet(doc, "retry-service chạy tại port 3003.")
    add_bullet(doc, "Input chính là topic send_failed.")
    add_bullet(doc, "Output là topic send_retry hoặc dead_letter.")
    add_bullet(doc, "Sử dụng retry theo exponential backoff.")
    add_info_table(
        doc,
        [
            ("Port", "3003"),
            ("Input topic", "send_failed"),
            ("Output topic", "send_retry, dead_letter"),
            ("Vai trò", "Retry command lỗi và quản lý dead letter"),
        ],
    )
    add_figure_box(doc, "Hình 13. Ảnh retry-service đang chạy")
    add_figure_box(doc, "Hình 14. Ảnh endpoint status của retry-service")
    add_figure_box(doc, "Hình 15. Ảnh topic send_failed")
    add_figure_box(doc, "Hình 16. Ảnh topic dead_letter")

    add_section(doc, "5. Luồng xử lý End-to-end")
    add_para(
        doc,
        "Toàn bộ hệ thống được xây dựng theo kiến trúc event-driven. "
        "Người dùng bình luận vào bài viết trên Facebook Page, Facebook gửi event về webhook-service, "
        "sau đó event được đưa vào Kafka để các service phía sau xử lý bất đồng bộ. "
        "Luồng này giúp từng service tách biệt trách nhiệm, dễ debug và dễ mở rộng hơn so với cách xử lý đồng bộ trong một service duy nhất."
    )

    add_subsection(doc, "5.1. Luồng thành công")
    add_bullet(doc, "Bước 1: Người dùng bình luận vào bài viết trên Facebook Page.")
    add_bullet(doc, "Bước 2: Facebook gửi event về endpoint /webhook của webhook-service.")
    add_bullet(doc, "Bước 3: webhook-service parse dữ liệu và publish vào topic raw_events.")
    add_bullet(doc, "Bước 4: core-service consume raw_events, phân tích nội dung và sinh command.")
    add_bullet(doc, "Bước 5: Command được publish sang reply_commands.")
    add_bullet(doc, "Bước 6: backend-api consume reply_commands và gọi Facebook Graph API để phản hồi.")
    add_bullet(doc, "Bước 7: Kết quả phản hồi xuất hiện trực tiếp dưới bài viết của page.")
    add_figure_box(doc, "Hình 17. Ảnh luồng thành công từ comment thật đến phản hồi trên page")

    add_subsection(doc, "5.2. Luồng phát sinh lỗi")
    add_bullet(doc, "Bước 1: backend-api nhận command nhưng gọi Facebook Graph API thất bại.")
    add_bullet(doc, "Bước 2: Command lỗi được đưa sang topic send_failed.")
    add_bullet(doc, "Bước 3: retry-service consume send_failed và retry theo backoff.")
    add_bullet(doc, "Bước 4: Nếu retry thành công, command quay lại send_retry để backend-api xử lý lại.")
    add_bullet(doc, "Bước 5: Nếu vượt quá số lần retry, command được đưa vào dead_letter.")
    add_figure_box(doc, "Hình 18. Ảnh luồng lỗi với send_failed, send_retry và dead_letter")

    add_section(doc, "6. Các lỗi gặp phải và cách xử lý")
    add_para(
        doc,
        "Trong quá trình làm bài, em gặp một số lỗi quan trọng liên quan tới cấu hình và tích hợp hệ thống. "
        "Việc ghi lại các lỗi này giúp làm rõ quá trình debug và chứng minh em đã hiểu được nguyên nhân của từng vấn đề."
    )
    add_bullet(doc, "Lỗi verify webhook thất bại do callback URL sai hoặc service chưa chạy đúng route /webhook.")
    add_bullet(doc, "Lỗi Kafka broker chưa sẵn sàng làm các service .NET không connect được tới localhost:9092.")
    add_bullet(doc, "Lỗi SQL Server local chiếm port 1433 nên phải đổi SQL container sang host port 1435.")
    add_bullet(doc, "Lỗi Page Access Token hết hạn làm backend-api gửi Facebook Graph API bị thất bại.")
    add_bullet(doc, "Lỗi self-loop khi page tự phản hồi comment của mình và hệ thống lại nhận chính phản hồi đó.")
    add_para(
        doc,
        "Sau khi phân tích, em đã xử lý bằng cách cấu hình lại callback URL, restart các service sau khi Kafka sẵn sàng, "
        "đổi port SQL Server container, tạo lại Page Access Token mới và bổ sung logic bỏ qua event do chính Page sinh ra."
    )
    add_figure_box(doc, "Hình 19. Ảnh một lỗi tiêu biểu và log xử lý", "Có thể dùng ảnh retry-service, backend-api hoặc Kafka reconnect.")

    add_section(doc, "7. Kết luận")
    add_para(
        doc,
        "Sau quá trình thực hiện, em đã xây dựng được hệ thống tích hợp Facebook Page Webhook với Kafka và các microservice .NET theo đúng hướng event-driven. "
        "Hệ thống có khả năng tiếp nhận sự kiện thật từ Facebook, đẩy dữ liệu qua Kafka, phân tích nội dung, sinh command và phản hồi lại trên page. "
        "Bên cạnh đó, hệ thống cũng có cơ chế retry, dead letter và lưu vết lỗi trong database để đảm bảo khả năng theo dõi và xử lý sự cố."
    )
    add_para(
        doc,
        "Qua bài tập này, em hiểu rõ hơn về cách cấu hình Meta Webhooks, cách tổ chức nhiều service giao tiếp qua Kafka, "
        "cách debug các lỗi tích hợp thực tế như token hết hạn, broker không sẵn sàng, route không khớp, và loop do event tự sinh ra. "
        "Đây là những kinh nghiệm quan trọng khi xây dựng hệ thống tích hợp với nền tảng bên ngoài."
    )

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
