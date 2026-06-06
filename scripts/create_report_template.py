from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
OUTPUT = ROOT / "Bao_cao_mau_qua_trinh_lam_bai_Facebook_Page.docx"

ACCENT = RGBColor(46, 116, 181)
ACCENT_DARK = RGBColor(31, 77, 120)
GRAY = RGBColor(90, 90, 90)
LIGHT_FILL = "F2F4F7"
BOX_FILL = "F7F7F7"


def set_document_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)

    for style_name, size, color in [
        ("Title", 22, RGBColor(0, 0, 0)),
        ("Heading 1", 16, ACCENT),
        ("Heading 2", 13, ACCENT),
        ("Heading 3", 12, ACCENT_DARK),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        if style_name != "Title":
            style.font.bold = True


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT, color=None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(10.5)
    if color is not None:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_spacer(doc: Document, points: int) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(points)


def add_cover(doc: Document) -> None:
    add_spacer(doc, 24)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BÁO CÁO QUÁ TRÌNH THỰC HIỆN BÀI TẬP")
    run.bold = True
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(22)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Tích hợp Facebook Page Webhook - Kafka - Microservices .NET")
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(15)
    run.font.color.rgb = GRAY

    add_spacer(doc, 18)

    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.columns[0].width = Inches(2)
    table.columns[1].width = Inches(4.2)
    rows = [
        ("Môn học / Học phần", "[Điền tên môn học]"),
        ("Đề tài", "Xây dựng hệ thống xử lý Facebook Page bằng Webhook và Kafka"),
        ("Sinh viên thực hiện", "[Họ và tên]"),
        ("MSSV", "[Mã số sinh viên]"),
        ("Lớp", "[Tên lớp]"),
        ("Ngày nộp", "[dd/mm/yyyy]"),
    ]
    for index, (label, value) in enumerate(rows):
        set_cell_shading(table.cell(index, 0), LIGHT_FILL)
        set_cell_text(table.cell(index, 0), label, bold=True)
        set_cell_text(table.cell(index, 1), value)

    add_spacer(doc, 24)
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("Gợi ý: thay ảnh bìa bằng logo trường, ảnh chụp giao diện Kafka UI hoặc sơ đồ hệ thống nếu muốn báo cáo đẹp hơn.")
    run.italic = True
    run.font.size = Pt(10)
    run.font.name = "Arial"

    doc.add_page_break()


def add_section_title(doc: Document, text: str, level: int = 1) -> None:
    doc.add_paragraph(text, style=f"Heading {level}")


def add_body(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.name = "Arial"
        r1._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        r1._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    r2 = p.add_run(text)
    r2.font.name = "Arial"
    r2._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    r2._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(item)
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(item)
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")


def add_info_table(doc: Document, rows: list[tuple[str, str]], col_widths: tuple[float, float] = (2.2, 4.3)) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Inches(col_widths[0])
    table.columns[1].width = Inches(col_widths[1])
    for index, (left, right) in enumerate(rows):
        set_cell_shading(table.cell(index, 0), LIGHT_FILL)
        set_cell_text(table.cell(index, 0), left, bold=True)
        set_cell_text(table.cell(index, 1), right)
    add_spacer(doc, 8)


def add_image_placeholder(doc: Document, title: str, description: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, BOX_FILL)
    cell.width = Inches(6.2)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f"{title}\n")
    r1.bold = True
    r1.font.name = "Arial"
    r1.font.size = Pt(11)
    r2 = p.add_run(description + "\n")
    r2.font.name = "Arial"
    r2.font.size = Pt(10)
    r3 = p.add_run("[Chèn ảnh minh chứng tại đây]")
    r3.italic = True
    r3.font.name = "Arial"
    r3.font.size = Pt(10)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    add_spacer(doc, 4)
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run(title)
    run.italic = True
    run.font.name = "Arial"
    run.font.size = Pt(10)
    add_spacer(doc, 8)


def add_flow_box(doc: Document, lines: list[str]) -> None:
    table = doc.add_table(rows=len(lines), cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, line in enumerate(lines):
        set_cell_shading(table.cell(i, 0), BOX_FILL if i % 2 == 0 else "FFFFFF")
        set_cell_text(table.cell(i, 0), line, bold=i in {0, len(lines) - 1}, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_spacer(doc, 8)


def add_footer(section, label: str) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(label + " | ")
    run.font.name = "Arial"
    run.font.size = Pt(9)
    page = p.add_run("Trang ")
    page.font.name = "Arial"
    page.font.size = Pt(9)

    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    p._p.append(fld)


def build_document() -> Path:
    doc = Document()
    set_document_defaults(doc)
    add_footer(doc.sections[0], "Báo cáo quá trình làm bài")

    add_cover(doc)

    doc.add_paragraph("MỤC LỤC GỢI Ý", style="Heading 1")
    add_numbered(
        doc,
        [
            "Thông tin bài toán và mục tiêu thực hiện",
            "Môi trường, công nghệ và công cụ sử dụng",
            "Kiến trúc hệ thống và giải thích luồng xử lý",
            "Quá trình triển khai theo từng giai đoạn",
            "Cấu hình Facebook Developer, ngrok, Docker, SQL Server và Kafka",
            "Kiểm thử, ảnh minh chứng và kết quả đạt được",
            "Các lỗi gặp phải và cách xử lý",
            "Đánh giá kết quả, hạn chế và hướng phát triển",
        ],
    )

    add_section_title(doc, "1. Thông tin bài toán và mục tiêu thực hiện", 1)
    add_body(doc, "Phần này dùng để giới thiệu ngắn gọn đề tài, bối cảnh bài tập và yêu cầu đầu ra của hệ thống. Có thể viết theo dạng 1 đến 2 đoạn ngắn, tránh lan man.")
    add_body(doc, "Mẫu nội dung tham khảo: Bài tập yêu cầu xây dựng một hệ thống tiếp nhận sự kiện từ Facebook Page thông qua Webhook, đưa dữ liệu vào Kafka để xử lý bất đồng bộ, sau đó phân loại nội dung bình luận hoặc tin nhắn, sinh lệnh phản hồi và gọi Facebook Graph API để trả lời hoặc ẩn bình luận.", "Gợi ý viết: ")
    add_bullets(
        doc,
        [
            "Mục tiêu 1: Xây dựng được hệ thống microservices chạy local bằng .NET.",
            "Mục tiêu 2: Tích hợp Facebook Page Webhook để nhận sự kiện thật.",
            "Mục tiêu 3: Sử dụng Kafka để truyền sự kiện giữa các service.",
            "Mục tiêu 4: Có cơ chế retry, dead letter, idempotency và lưu vết xử lý.",
        ],
    )

    add_section_title(doc, "2. Môi trường, công nghệ và công cụ sử dụng", 1)
    add_info_table(
        doc,
        [
            ("Ngôn ngữ / Framework", ".NET / ASP.NET Core"),
            ("Message Broker", "Apache Kafka"),
            ("Cơ sở dữ liệu", "SQL Server"),
            ("Hạ tầng local", "Docker Compose"),
            ("Expose Webhook", "ngrok"),
            ("Nguồn sự kiện", "Facebook Page Webhook"),
            ("UI kiểm tra Kafka", "Kafka UI tại http://localhost:8085"),
        ],
    )
    add_image_placeholder(
        doc,
        "Hình 1. Ảnh cấu trúc thư mục dự án",
        "Chụp cây thư mục project gồm webhook-service, core-service, backend-api, retry-service, shared-models."
    )

    add_section_title(doc, "3. Kiến trúc hệ thống và giải thích luồng xử lý", 1)
    add_body(doc, "Phần này là phần quan trọng nhất của báo cáo. Nên giải thích rõ vai trò từng service và luồng dữ liệu đi như thế nào từ Facebook vào hệ thống rồi quay ngược lại Facebook.")
    add_flow_box(
        doc,
        [
            "Facebook Page / User Comment",
            "Webhook -> webhook-service (port 3001)",
            "Publish event -> Kafka topic raw_events",
            "Consume -> core-service (port 3002)",
            "Phân loại nội dung, phát sinh command",
            "Publish -> Kafka topic reply_commands / send_retry",
            "Consume -> backend-api (port 3000)",
            "Gọi Facebook Graph API để reply hoặc hide comment",
            "Nếu lỗi -> send_failed -> retry-service (port 3003)",
            "Retry nhiều lần, quá ngưỡng -> dead_letter",
        ],
    )
    add_info_table(
        doc,
        [
            ("webhook-service", "Nhận HTTP webhook từ Facebook, kiểm tra verify token, publish sự kiện vào raw_events."),
            ("core-service", "Consume raw_events, phát hiện spam / toxic / negative, tạo command xử lý."),
            ("backend-api", "Consume reply_commands và send_retry, gọi Facebook Graph API, lưu failed command."),
            ("retry-service", "Consume send_failed, retry theo exponential backoff, đẩy sang dead_letter nếu vượt ngưỡng."),
        ],
    )
    add_image_placeholder(
        doc,
        "Hình 2. Sơ đồ kiến trúc hệ thống",
        "Có thể chụp ảnh tự vẽ trên draw.io hoặc dùng sơ đồ chữ trong báo cáo rồi chụp lại."
    )

    add_section_title(doc, "4. Quá trình triển khai theo từng giai đoạn", 1)
    add_section_title(doc, "4.1. Chuẩn bị source code và môi trường", 2)
    add_numbered(
        doc,
        [
            "Clone hoặc pull source code về máy.",
            "Kiểm tra các service và cấu hình còn thiếu.",
            "Bổ sung file project, appsettings và các dependency cần thiết.",
            "Khởi động Docker để chạy Kafka, Zookeeper, SQL Server và Kafka UI.",
        ],
    )
    add_image_placeholder(
        doc,
        "Hình 3. Docker Compose và danh sách container đang chạy",
        "Chụp lệnh docker-compose up -d và docker ps."
    )

    add_section_title(doc, "4.2. Cấu hình Facebook Developer và webhook", 2)
    add_bullets(
        doc,
        [
            "Tạo app trên Meta for Developers.",
            "Bật Webhooks cho object Page.",
            "Thiết lập Callback URL trỏ tới URL ngrok + /webhook.",
            "Thiết lập Verify Token trùng với giá trị trong appsettings.",
            "Subscribe field feed để nhận sự kiện comment trên bài viết.",
            "Subscribe app vào đúng Page cần theo dõi.",
        ],
    )
    add_info_table(
        doc,
        [
            ("Callback URL mẫu", "https://<ngrok-domain>/webhook"),
            ("Verify Token", "[Điền verify token đang dùng]"),
            ("Page ID", "[Điền page id]"),
            ("App ID", "[Điền app id]"),
        ],
    )
    add_image_placeholder(
        doc,
        "Hình 4. Màn hình cấu hình Webhook trên Facebook Developer",
        "Chụp phần Callback URL, Verify Token, subscribed fields và trạng thái verify thành công."
    )

    add_section_title(doc, "4.3. Chạy các service", 2)
    add_body(doc, "Liệt kê đúng thứ tự chạy để người đọc có thể chạy lại hệ thống theo báo cáo.")
    add_numbered(
        doc,
        [
            "Chạy Docker Compose để lên Kafka, Zookeeper, SQL Server, Kafka UI.",
            "Chạy webhook-service.",
            "Chạy core-service.",
            "Chạy backend-api.",
            "Chạy retry-service.",
            "Chạy ngrok để public port 3001.",
        ],
    )
    add_image_placeholder(
        doc,
        "Hình 5. Các cửa sổ terminal khi chạy 4 service",
        "Chụp các terminal thể hiện service listen ở port 3000, 3001, 3002, 3003."
    )

    add_section_title(doc, "5. Cấu hình chi tiết quan trọng", 1)
    add_info_table(
        doc,
        [
            ("webhook-service", "Port 3001, route /webhook, verify token, page access token."),
            ("core-service", "Consume raw_events, publish reply_commands."),
            ("backend-api", "Port 3000, SQL Server localhost:1435, gọi Graph API."),
            ("retry-service", "Port 3003, xử lý send_failed, send_retry, dead_letter."),
            ("Kafka", "Bootstrap server localhost:9092."),
            ("Kafka UI", "http://localhost:8085."),
            ("SQL Server", "localhost:1435, database ApiFacebookDb."),
        ],
    )
    add_body(doc, "Nếu giảng viên yêu cầu, có thể bổ sung 1 bảng riêng cho appsettings.json của từng service và giải thích ý nghĩa từng tham số.")

    add_section_title(doc, "6. Kiểm thử và ảnh minh chứng", 1)
    add_body(doc, "Nên chia rõ giữa kiểm thử kỹ thuật và kiểm thử nghiệp vụ. Mỗi test case nên có ảnh minh chứng và kết quả mong đợi.")
    add_info_table(
        doc,
        [
            ("Test case 1", "Verify webhook thành công."),
            ("Test case 2", "Nhận event thật từ Facebook comment vào raw_events."),
            ("Test case 3", "core-service sinh reply_commands."),
            ("Test case 4", "backend-api reply comment hoặc hide comment thành công."),
            ("Test case 5", "Lỗi Graph API -> send_failed -> retry-service -> dead_letter."),
        ],
    )
    add_image_placeholder(
        doc,
        "Hình 6. Topic raw_events có dữ liệu",
        "Chụp Kafka UI khi raw_events nhận sự kiện từ mock hoặc từ Facebook thật."
    )
    add_image_placeholder(
        doc,
        "Hình 7. Topic reply_commands có dữ liệu",
        "Chụp Kafka UI khi core-service đã tạo command phản hồi."
    )
    add_image_placeholder(
        doc,
        "Hình 8. Kết quả phản hồi thực tế trên Facebook Page",
        "Chụp bình luận của user và phản hồi hoặc ẩn bình luận của Page."
    )

    add_section_title(doc, "7. Các lỗi gặp phải và cách xử lý", 1)
    add_body(doc, "Phần này nên viết thành bảng hoặc theo từng ý ngắn. Giảng viên thường đánh giá cao phần mô tả lỗi và cách tự xử lý.")
    add_info_table(
        doc,
        [
            ("Lỗi 1", "Webhook verify thất bại do URL sai hoặc service chưa chạy."),
            ("Cách xử lý", "Kiểm tra route /webhook, verify token và ngrok forward đúng port 3001."),
            ("Lỗi 2", "Kafka broker down hoặc service chạy trước khi Kafka sẵn sàng."),
            ("Cách xử lý", "Khởi động lại Docker Compose, kiểm tra docker ps và restart các service .NET."),
            ("Lỗi 3", "Page Access Token hết hạn làm Graph API trả lỗi."),
            ("Cách xử lý", "Tạo token mới, cập nhật appsettings và chạy lại backend-api."),
            ("Lỗi 4", "Hệ thống tự loop khi Page phản hồi lại chính comment của mình."),
            ("Cách xử lý", "Chặn event do Page tự tạo, bỏ qua event có SenderId trùng PageId."),
        ],
    )
    add_image_placeholder(
        doc,
        "Hình 9. Một lỗi tiêu biểu và log xử lý",
        "Có thể chụp log retry-service, backend-api hoặc Kafka reconnect để minh họa cách debug."
    )

    add_section_title(doc, "8. Kết quả đạt được", 1)
    add_bullets(
        doc,
        [
            "Nhận được sự kiện thật từ Facebook Page thông qua webhook.",
            "Đưa sự kiện vào Kafka và xử lý theo kiến trúc bất đồng bộ.",
            "Phân luồng được spam, negative hoặc yêu cầu phản hồi thông thường.",
            "Thực hiện được reply comment hoặc hide comment qua backend-api.",
            "Có cơ chế retry và dead letter cho các lệnh lỗi.",
        ],
    )

    add_section_title(doc, "9. Hạn chế và hướng phát triển", 1)
    add_bullets(
        doc,
        [
            "Hệ thống hiện chủ yếu chạy local, chưa triển khai production.",
            "Việc phân loại spam / toxic còn dựa nhiều vào rule-based hoặc fallback đơn giản.",
            "Token Facebook có thể hết hạn, cần cơ chế quản lý credential tốt hơn.",
            "Có thể mở rộng dashboard giám sát để xem trạng thái command và dead letter trực quan hơn.",
        ],
    )

    add_section_title(doc, "10. Kết luận", 1)
    add_body(doc, "Phần kết luận nên tóm tắt ngắn gọn những gì đã hoàn thành, những gì học được khi làm bài và mức độ đáp ứng yêu cầu đề bài.")
    add_body(doc, "Mẫu kết luận tham khảo: Sau quá trình triển khai, em đã xây dựng được hệ thống tích hợp Facebook Page Webhook với Kafka và các microservice .NET. Hệ thống có thể tiếp nhận sự kiện, xử lý bất đồng bộ, sinh phản hồi và có cơ chế retry khi xảy ra lỗi. Trong quá trình thực hiện, em cũng rút ra kinh nghiệm về cấu hình webhook, quản lý token, xử lý lỗi Kafka và đồng bộ giữa các service.")

    add_section_title(doc, "Phụ lục A. Checklist ảnh minh chứng nên có", 1)
    add_bullets(
        doc,
        [
            "Ảnh docker ps hoặc Docker Desktop đang chạy đủ container.",
            "Ảnh terminal của 4 service listen đúng port.",
            "Ảnh cấu hình Webhook trên Meta for Developers.",
            "Ảnh verify webhook thành công.",
            "Ảnh Kafka UI có topic raw_events, reply_commands, send_failed, send_retry, dead_letter.",
            "Ảnh comment thật trên Facebook và phản hồi của hệ thống.",
            "Ảnh một lỗi và cách bạn xử lý được lỗi đó.",
        ],
    )

    add_section_title(doc, "Phụ lục B. Dàn ý ngắn để thuyết trình báo cáo", 1)
    add_numbered(
        doc,
        [
            "Giới thiệu bài toán và mục tiêu.",
            "Giải thích kiến trúc 4 service và Kafka.",
            "Trình bày quy trình cấu hình Facebook Webhook.",
            "Demo luồng chạy thật với Kafka UI và Facebook Page.",
            "Nêu các lỗi tiêu biểu và cách xử lý.",
            "Kết luận và hướng phát triển.",
        ],
    )

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    result = build_document()
    print(result)
