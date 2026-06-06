from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
OUTPUT = ROOT / "Mau_bao_cao_co_ban_Facebook_Page_API.docx"

ACCENT = RGBColor(46, 116, 181)
GRAY = RGBColor(90, 90, 90)
BOX_FILL = "F7F7F7"


def set_font(run, size=11, bold=False, italic=False, color=None):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def set_document_defaults(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)

    for name, size in [("Heading 1", 14), ("Heading 2", 12)]:
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = ACCENT


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_para(doc: Document, text: str, bold=False, italic=False, center=False, color=None, after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, bold=bold, italic=italic, color=color)
    return p


def add_section(doc: Document, title: str):
    doc.add_paragraph(title, style="Heading 1")


def add_subsection(doc: Document, title: str):
    doc.add_paragraph(title, style="Heading 2")


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_font(run)


def add_figure_box(doc: Document, title: str, note: str = ""):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, BOX_FILL)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    r1 = p.add_run(title + "\n")
    set_font(r1, bold=True)
    r2 = p.add_run("[Chèn ảnh minh chứng tại đây]")
    set_font(r2, italic=True, color=GRAY)
    if note:
        r3 = p.add_run("\n" + note)
        set_font(r3, size=10, color=GRAY)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(title)
    set_font(r, size=10, italic=True)
    cap.paragraph_format.space_after = Pt(6)


def add_name_block(doc: Document):
    add_para(doc, "Họ và tên: [Điền họ và tên]", bold=True, after=2)
    add_para(doc, "MSSV: [Điền mã số sinh viên]", bold=True, after=10)


def build():
    doc = Document()
    set_document_defaults(doc)

    add_name_block(doc)

    add_section(doc, "1. Webhook")
    add_para(doc, "Mô tả ngắn: phần này trình bày việc cấu hình webhook, verify callback URL và quá trình nhận sự kiện từ Facebook Page.")
    add_bullet(doc, "webhook-service chạy trên port 3001.")
    add_bullet(doc, "Facebook Developer dùng Callback URL trỏ về ngrok + /webhook.")
    add_bullet(doc, "Verify Token phải trùng với cấu hình trong appsettings.")
    add_figure_box(doc, "Hình 1. Ảnh webhook-service chạy trên port 3001")
    add_figure_box(doc, "Hình 2. Ảnh cấu hình Webhook trên Facebook Developer")
    add_figure_box(doc, "Hình 3. Ảnh comment thật trên bài viết của page")
    add_figure_box(doc, "Hình 4. Ảnh Kafka UI nhận raw_events")

    add_section(doc, "2. Core Service")
    add_para(doc, "Mô tả ngắn: phần này trình bày cách core-service consume raw_events, phân tích nội dung và sinh command phản hồi.")
    add_bullet(doc, "core-service chạy trên port 3002.")
    add_bullet(doc, "Kiểm tra trường hợp bình luận tích cực, hỏi giá, khiếu nại, spam.")
    add_bullet(doc, "Kết quả xử lý sẽ được publish sang reply_commands.")
    add_figure_box(doc, "Hình 5. Ảnh core-service đang chạy")
    add_figure_box(doc, "Hình 6. Ảnh bình luận tích cực hoặc hỏi giá")
    add_figure_box(doc, "Hình 7. Ảnh Kafka UI sau khi core-service xử lý")
    add_figure_box(doc, "Hình 8. Ảnh trường hợp spam hoặc toxic comment")

    add_section(doc, "3. Backend API")
    add_para(doc, "Mô tả ngắn: backend-api là service duy nhất được gọi Facebook Graph API để trả lời bình luận hoặc ẩn bình luận.")
    add_bullet(doc, "backend-api chạy trên port 3000.")
    add_bullet(doc, "Consume các topic reply_commands và send_retry.")
    add_bullet(doc, "Lưu vết failed_commands và idempotency key trong database.")
    add_figure_box(doc, "Hình 9. Ảnh backend-api đang chạy")
    add_figure_box(doc, "Hình 10. Ảnh health endpoint hoặc log consume command")
    add_figure_box(doc, "Hình 11. Ảnh topic reply_commands")
    add_figure_box(doc, "Hình 12. Ảnh bảng dữ liệu trong SQL Server hoặc failed_commands")

    add_section(doc, "4. Retry Service")
    add_para(doc, "Mô tả ngắn: retry-service xử lý các command bị lỗi, retry lại theo số lần cấu hình, sau đó đưa sang dead_letter nếu vượt ngưỡng.")
    add_bullet(doc, "retry-service chạy trên port 3003.")
    add_bullet(doc, "Input chính là topic send_failed.")
    add_bullet(doc, "Nếu retry vượt quá ngưỡng thì message sẽ vào dead_letter.")
    add_figure_box(doc, "Hình 13. Ảnh retry-service đang chạy")
    add_figure_box(doc, "Hình 14. Ảnh endpoint status của retry-service")
    add_figure_box(doc, "Hình 15. Ảnh topic send_failed")
    add_figure_box(doc, "Hình 16. Ảnh topic dead_letter")

    add_section(doc, "5. Luồng xử lý End-to-end")
    add_para(doc, "Phần này mô tả lại toàn bộ luồng chạy từ lúc người dùng bình luận trên Facebook đến khi hệ thống trả lời hoặc xử lý lỗi.")

    add_subsection(doc, "5.1. Luồng thành công")
    add_bullet(doc, "Người dùng bình luận vào bài viết trên page.")
    add_bullet(doc, "webhook-service nhận event và đẩy vào raw_events.")
    add_bullet(doc, "core-service phân tích nội dung và publish sang reply_commands.")
    add_bullet(doc, "backend-api nhận command và gọi Facebook Graph API để phản hồi.")
    add_figure_box(doc, "Hình 17. Ảnh luồng thành công - từ comment đến phản hồi thật")

    add_subsection(doc, "5.2. Luồng phát sinh lỗi")
    add_bullet(doc, "backend-api gọi Facebook Graph API thất bại do token hết hạn hoặc comment id lỗi.")
    add_bullet(doc, "command bị đẩy sang send_failed.")
    add_bullet(doc, "retry-service retry theo backoff và đưa vào dead_letter nếu vượt ngưỡng.")
    add_figure_box(doc, "Hình 18. Ảnh luồng lỗi - send_failed, send_retry, dead_letter")

    add_section(doc, "6. Các lỗi gặp phải và cách xử lý")
    add_bullet(doc, "Lỗi verify webhook do sai URL hoặc verify token.")
    add_bullet(doc, "Lỗi Kafka chưa sẵn sàng nên service không connect được.")
    add_bullet(doc, "Lỗi Page Access Token hết hạn dẫn đến Graph API trả lỗi.")
    add_bullet(doc, "Lỗi loop khi page tự phản hồi chính comment của mình.")
    add_para(doc, "Ở mỗi lỗi, bạn nên viết ngắn gọn: nguyên nhân, biểu hiện, và cách xử lý đã làm.")
    add_figure_box(doc, "Hình 19. Ảnh một lỗi tiêu biểu và cách bạn xử lý")

    add_section(doc, "7. Kết luận")
    add_para(doc, "Phần này viết khoảng 1 đoạn ngắn để tổng kết những gì đã hoàn thành, những kỹ năng học được và đánh giá mức độ đáp ứng yêu cầu bài tập.")
    add_para(doc, "Gợi ý: Sau quá trình thực hiện, em đã xây dựng được hệ thống tích hợp Facebook Page Webhook với Kafka và các microservice .NET. Hệ thống có thể tiếp nhận sự kiện, xử lý bất đồng bộ, sinh phản hồi và có cơ chế retry khi xảy ra lỗi. Qua bài tập, em hiểu rõ hơn về cách cấu hình webhook, vận hành Kafka và debug hệ thống nhiều service.")

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
